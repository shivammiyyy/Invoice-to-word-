import pytesseract
from PIL import Image
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import streamlit as st
import os

# Set up Tesseract executable path (adjust this to your Tesseract-OCR installation path)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def extract_text(image_path):
    return pytesseract.image_to_string(Image.open(image_path))

def create_word_doc(text, output_path):
    doc = Document()

    # Set font for the whole document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier New'  # Using a monospaced font to preserve layout better
    font.size = Pt(10)

    # Add extracted text to the document
    lines = text.split('\n')
    for line in lines:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
    
    doc.save(output_path)

def main():
    st.title("Invoice Reader")

    uploaded_file = st.file_uploader("Choose an image file", type=["tif", "jpg", "jpeg", "png", "gif"])
    
    if uploaded_file is not None:
        # Save the uploaded file temporarily
        with open(os.path.join("temp", uploaded_file.name), "wb") as f:
            f.write(uploaded_file.getbuffer())

        image_path = os.path.join("temp", uploaded_file.name)

        # Extract text from image
        extracted_text = extract_text(image_path)
        
        # Display extracted text
        st.subheader("Extracted Text")
        st.text(extracted_text)
        
        # Create Word document
        output_path = os.path.join("output", "extracted_text.docx")
        create_word_doc(extracted_text, output_path)
        
        # Provide a link to download the Word document
        with open(output_path, "rb") as f:
            st.download_button("Download Word Document", f, file_name="extracted_text.docx")

if __name__ == "__main__":
    os.makedirs("temp", exist_ok=True)
    os.makedirs("output", exist_ok=True)
    main()
